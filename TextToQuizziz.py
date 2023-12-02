import docx
import openpyxl

class Question: 
    def __init__(self, question):
        self.question = question
        self.answer = []
        self.rightAns = -1

    def addAnswer(self, answer):
        self.answer.append(answer)

    def setRightAns(self, rightAns):
        self.rightAns = rightAns

    def getRightAns(self):
        return self.rightAns

    def getAns(self):
        return self.answer

    def getAnsLen(self):
        return len(self.answer)
    
    def getQuest(self):
        return self.question
    
    def __str__(self):
        return "Question: " + self.question + "\nAnswer: " + str(self.answer) + "\nRight Answer: " + str(self.rightAns) + "\n"

class Converter:
    def __init__(self, docfile, xlsxfile):
        self.question = []

        self.docfile = docfile
        self.xlsxfile = xlsxfile

        self.docf = docx.Document(self.docfile)
        self.workbook = openpyxl.load_workbook(self.xlsxfile)
        self.worksheet = self.workbook["Create a Quiz"]

    def readtxt(self):
        currentLine = 0
        numOfLine = len(self.docf.paragraphs)

        while(currentLine < numOfLine):
            quest = Question(self.docf.paragraphs[currentLine].text.strip())
            if (quest.getQuest() == ""):
                currentLine += 1
                continue

            numOfAns = 4
            if ("Nhận định trên đúng hay sai" in quest.getQuest()):
                numOfAns = 2

            rightAns = -1
            index = 1
            while (quest.getAnsLen() < numOfAns):
                tmpAns = self.docf.paragraphs[currentLine + index]
                index += 1

                quest.addAnswer(tmpAns.text.strip())

                for run in tmpAns.runs:
                    if (str(run.font.color.rgb) == "FF0000"):
                        rightAns = index - 1
                        quest.setRightAns(rightAns)

            print(quest)
            currentLine += numOfAns + 1
            self.question.append(quest)

    def addSheet(self):
        row_num = 3

        for quest in self.question:
            self.worksheet.cell(row=row_num, column=1).value = quest.getQuest()
            self.worksheet.cell(row=row_num, column=2).value = 'Multiple Choice'

            ansArr = quest.getAns()
            numOfAns = len(ansArr)
            for i in range(numOfAns):
                self.worksheet.cell(row=row_num, column=3 + i).value = ansArr[i]

            self.worksheet.cell(row=row_num, column=7).value = quest.getRightAns()
            self.worksheet.cell(row=row_num, column=8).value = 60
            row_num += 1

        self.workbook.save(self.xlsxfile) 
        print("done")

c = Converter("H:\\Download\\question.docx", "H:\\Download\\question.xlsx")
c.readtxt()
c.addSheet()